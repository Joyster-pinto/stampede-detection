"""
generate_word_report.py
========================
Compiles all markdown reports and images from the 7 experiments into a single 
beautiful Microsoft Word (.docx) document for final submission.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    # Make headings black instead of default blue
    for run in heading.runs:
        run.font.color.rgb = None

def parse_markdown_to_docx(doc, md_filepath, img_dir):
    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle Tables
        if line.startswith('|'):
            if '---' in line:
                continue # Skip separator
            row_data = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(row_data)
            in_table = True
            continue
        elif in_table:
            # We finished a table block
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_data):
                    row = table.rows[i].cells
                    for j, cell_text in enumerate(row_data):
                        # Clean up markdown bold from cell text
                        cell_text = cell_text.replace('**', '')
                        row[j].text = cell_text
                        if i == 0: # Header row bold
                            for paragraph in row[j].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                table_data = []
            in_table = False
            
        # Handle Headings
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
            
        # Handle List Items
        elif line.startswith('- '):
            text = line[2:].replace('**', '') # Strip bold markdown for simplicity
            doc.add_paragraph(text, style='List Bullet')
            
        # Handle Images referenced in markdown or just general text
        else:
            # Strip markdown bolding
            text = line.replace('**', '')
            
            # Simple check for image markdown syntax (if they exist in the md)
            if line.startswith('![') and '](' in line:
                # We won't parse it directly from md if it's too complex, 
                # but our reports mostly just list assets at the end.
                continue 
            else:
                doc.add_paragraph(text)

    # After reading the text, append all PNGs from that experiment folder
    for f in os.listdir(img_dir):
        if f.endswith('.png'):
            img_path = os.path.join(img_dir, f)
            try:
                # Add a caption
                caption = doc.add_paragraph()
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption.add_run(f"Figure: {f.replace('.png', '')}").italic = True
                
                # Add image
                doc.add_picture(img_path, width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1] 
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                doc.add_paragraph() # Spacing
            except Exception as e:
                print(f"Could not add image {img_path}: {e}")

def create_compiled_report():
    base_dir = r"D:\Research\Experiment Output"
    output_docx = r"D:\Research\Final_Submission_Report.docx"
    
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Stampede Detection using Dense Optical Flow and BiLSTM+Attention', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('\nFinal Research Internship Submission')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()
    
    # Define experiment folders in order
    experiments = [
        "Experiment_1_Model_Comparison",
        "Experiment_2_Feature_Ablation",
        "Experiment_3_CrossDataset",
        "Experiment_4_Early_Detection",
        "Experiment_5_Robustness",
        "Experiment_6_Cost_Analysis",
        "Experiment_7_Explainability"
    ]
    
    for exp_folder in experiments:
        exp_path = os.path.join(base_dir, exp_folder)
        if not os.path.isdir(exp_path):
            continue
            
        # Find the markdown report in this folder
        md_file = None
        for f in os.listdir(exp_path):
            if f.endswith('.md'):
                md_file = os.path.join(exp_path, f)
                break
                
        if md_file:
            print(f"Processing {exp_folder}...")
            parse_markdown_to_docx(doc, md_file, exp_path)
            doc.add_page_break()
            
    doc.save(output_docx)
    print(f"Successfully generated {output_docx}")

if __name__ == "__main__":
    create_compiled_report()
